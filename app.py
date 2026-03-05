import os
import shutil
from pathlib import Path
from typing import List, Tuple

import gradio as gr
from dotenv import load_dotenv

import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda
from langchain_core.runnables.passthrough import RunnableAssign
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ------------------------------------------------------------------------------
# Config
# ------------------------------------------------------------------------------
load_dotenv()

DATA_DIR = Path(os.getenv("DATA_DIR", "./data"))
UPLOAD_DIR = DATA_DIR / "uploads"
VS_DIR = DATA_DIR / "vectorstore"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
VS_DIR.mkdir(parents=True, exist_ok=True)

_vs_cache: FAISS | None = None
_retriever_cache = None


def _require_env():
    if not os.getenv("ZAI_API_KEY"):
        raise RuntimeError("Missing ZAI_API_KEY. Put it in your environment or in a .env file.")


# ------------------------------------------------------------------------------
# Model + Embeddings (Z.ai OpenAI-compatible)
# ------------------------------------------------------------------------------
def get_llm():
    _require_env()
    return ChatOpenAI(
        api_key=os.getenv("ZAI_API_KEY"),
        base_url="https://api.z.ai/api/coding/paas/v4/",
        model="glm-4.7",
        temperature=0.7,
    )


def get_embeddings():
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")


# ------------------------------------------------------------------------------
# Loaders
# ------------------------------------------------------------------------------
def read_pdf_text(path: str) -> str:
    parts = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
    except Exception:
        parts = []

    if parts:
        return "\n\n".join(parts)

    # Fallback to PyMuPDF
    doc = fitz.open(path)
    try:
        for page in doc:
            t = page.get_text("text") or ""
            if t.strip():
                parts.append(t)
    finally:
        doc.close()

    return "\n\n".join(parts)


def read_docx_text(path: str) -> str:
    d = DocxDocument(path)
    paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    return "\n".join(paras)


def load_document(file_path: str) -> List[Document]:
    p = Path(file_path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        text = read_pdf_text(str(p))
    elif ext == ".docx":
        text = read_docx_text(str(p))
    elif ext in (".txt", ".md"):
        text = p.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, TXT, or MD.")

    if not text.strip():
        raise ValueError("No text extracted from file.")

    return [Document(page_content=text, metadata={"source": p.name})]


# ------------------------------------------------------------------------------
# Chunking + Vectorstore
# ------------------------------------------------------------------------------
def chunk_documents(docs: List[Document], chunk_size=1200, chunk_overlap=150) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", ";", ",", " ", ""],
    )
    chunks = splitter.split_documents(docs)
    for i, c in enumerate(chunks):
        c.metadata = dict(c.metadata or {})
        c.metadata["chunk_id"] = i
    return chunks


def build_or_replace_vectorstore(chunks: List[Document]) -> Tuple[int, int]:
    global _vs_cache, _retriever_cache

    if VS_DIR.exists():
        shutil.rmtree(VS_DIR)
    VS_DIR.mkdir(parents=True, exist_ok=True)

    embeddings = get_embeddings()
    vs = FAISS.from_documents(chunks, embeddings)
    vs.save_local(str(VS_DIR))

    _vs_cache = vs
    _retriever_cache = vs.as_retriever(search_kwargs={"k": 4})

    total_chars = sum(len(d.page_content) for d in chunks)
    return total_chars, len(chunks)


def load_vectorstore() -> FAISS:
    global _vs_cache, _retriever_cache

    if _vs_cache is not None:
        return _vs_cache

    embeddings = get_embeddings()
    if not VS_DIR.exists() or not any(VS_DIR.iterdir()):
        raise RuntimeError("No vectorstore found. Upload & ingest a document first.")

    vs = FAISS.load_local(
        str(VS_DIR),
        embeddings,
        # FAISS deserialization requires this flag; safe here because we wrote
        # the store ourselves from trusted local files.
        allow_dangerous_deserialization=True,
    )
    _vs_cache = vs
    _retriever_cache = vs.as_retriever(search_kwargs={"k": 4})
    return vs


def get_retriever():
    """Return cached retriever, loading the vectorstore if needed."""
    global _retriever_cache
    if _retriever_cache is None:
        vs = load_vectorstore()  # also populates _retriever_cache
        return vs.as_retriever(search_kwargs={"k": 4})
    return _retriever_cache


# ------------------------------------------------------------------------------
# RAG Chain (with conversation history)
# ------------------------------------------------------------------------------
def make_rag_chain(retriever):
    llm = get_llm()

    # across turns. history is a list of (user_msg, assistant_msg) tuples.
    prompt = ChatPromptTemplate.from_messages([
        ("system",
         "You are a helpful assistant for understanding contracts and policies. "
         "Answer using ONLY the provided context. "
         "If the answer is not in the context, say: 'I don't know from the document.'"),
        ("human",
         "Conversation so far:\n{history}\n\n"
         "Question: {question}\n\n"
         "Context:\n{context}\n\n"
         "Answer briefly, then add a 'Sources:' section with bullet points in the form:\n"
         "- [source | chunk id]"),
    ])

    def format_docs(docs: List[Document]) -> str:
        parts = []
        for d in docs:
            src = d.metadata.get("source", "document")
            cid = d.metadata.get("chunk_id", "?")
            parts.append(f"[{src} | chunk {cid}]\n{d.page_content}")
        return "\n\n---\n\n".join(parts)

    def _extract_text(content) -> str:
        if content is None:
            return ""
        if isinstance(content, str):
            return content
        if isinstance(content, list):
            parts = []
            for item in content:
                if isinstance(item, dict):
                    if isinstance(item.get("text"), str):
                        parts.append(item["text"])
                    elif isinstance(item.get("content"), str):
                        parts.append(item["content"])
                    elif isinstance(item.get("value"), str):
                        parts.append(item["value"])
                    else:
                        parts.append(str(item))
                else:
                    parts.append(str(item))
            return "\n".join(p for p in parts if p)
        return str(content)

    def format_history(history: list) -> str:
        if not history:
            return "(none)"

        lines = []
        for turn in history:
            # Gradio legacy format: [(user, assistant), ...]
            if isinstance(turn, (list, tuple)) and len(turn) == 2:
                user_msg, assistant_msg = turn
                lines.append(f"User: {_extract_text(user_msg)}")
                lines.append(f"Assistant: {_extract_text(assistant_msg)}")
                continue

            # Gradio messages format: [{"role": "user"|"assistant", "content": ...}, ...]
            if isinstance(turn, dict):
                role = str(turn.get("role", "")).strip().lower()
                content = _extract_text(turn.get("content"))
                if role == "user":
                    lines.append(f"User: {content}")
                elif role == "assistant":
                    lines.append(f"Assistant: {content}")
                elif role:
                    lines.append(f"{role.title()}: {content}")
                else:
                    lines.append(_extract_text(turn))
                continue

            # Any unexpected shape should not break the app.
            lines.append(_extract_text(turn))

        return "\n".join(line for line in lines if line.strip()) or "(none)"

    chain = (
        RunnableAssign({
            "docs": lambda x: retriever.invoke(x["question"]),
            "history_str": lambda x: format_history(x.get("history", [])),
        })
        | RunnableAssign({
            "context": lambda x: format_docs(x["docs"]),
            "history": lambda x: x["history_str"],
        })
        | prompt
        | llm
        | StrOutputParser()
    )
    return chain


# ------------------------------------------------------------------------------
# Gradio UI
# ------------------------------------------------------------------------------
def ingest_ui(file_obj):
    if file_obj is None:
        return "Please upload a file."

    # Gradio 4+ returns a filepath string directly; older versions return
    # an object with a .name attribute. Handle both.
    src = Path(file_obj if isinstance(file_obj, str) else file_obj.name)
    dst = UPLOAD_DIR / src.name
    if src.resolve() != dst.resolve():
        shutil.copyfile(src, dst)

    try:
        docs = load_document(str(dst))
        chunks = chunk_documents(docs)
        total_chars, n_chunks = build_or_replace_vectorstore(chunks)
        return (
            f"  Ingested successfully!\n"
            f"  chunks={n_chunks:,}  total_chunk_chars={total_chars:,}\n"
            f"  stored at: {VS_DIR}"
        )
    except Exception as e:
        return f"  Ingestion failed: {e}"


def chat_stream(message, history):
    try:
        retriever = get_retriever()
    except RuntimeError:
        yield "Please upload a document and click **Ingest** first."
        return

    # FIX #6: Pass history into the chain
    rag = make_rag_chain(retriever)

    buffer = ""
    try:
        for token in rag.stream({"question": message, "history": history}):
            buffer += token
            yield buffer
    except Exception as e:
        yield f"Error generating response: {e}"


def build_ui():
    with gr.Blocks(title="Contract RAG Assistant") as demo:
        gr.Markdown(
            "# Contract RAG Assistant\n"
            "1) Upload a document (PDF, DOCX, TXT, or MD)\n"
            "2) Click **Ingest**\n"
            "3) Ask questions in the **Chat** tab\n"
        )

        with gr.Tab("Upload"):
            uploader = gr.File(file_types=[".pdf", ".docx", ".txt", ".md"], label="Upload")
            ingest_btn = gr.Button("Ingest")
            status = gr.Textbox(label="Status", lines=3)
            ingest_btn.click(ingest_ui, inputs=[uploader], outputs=[status])

        with gr.Tab("Chat"):
            gr.ChatInterface(fn=chat_stream)

    return demo


if __name__ == "__main__":
    build_ui().queue().launch(share=False, debug=True)
